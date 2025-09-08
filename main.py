import json
from schema import RELATION_TYPES
from rules import clean_output
from prompt_template import PROMPT_TEMPLATE
import os
from openai import OpenAI
from tqdm import tqdm   # ✅ 新增：进度条库
from docx import Document
import requests, os
from concurrent.futures import ThreadPoolExecutor, as_completed
from retriever.retrieve_context import retrieve_context


# os.environ["HTTP_PROXY"] = "http://127.0.0.1:7890"
# os.environ["HTTPS_PROXY"] = "http://127.0.0.1:7890"

# from openai import OpenAI


# class OpenAILLM:
#     def __init__(self):
#         self.client = OpenAI(api_key="""
#     def generate(self, prompt):
#         response = self.client.chat.completions.create(
#             model="gpt-4o-mini",   # 也可以换成 "gpt-3.5-turbo"
#             messages=[
#                 {"role": "system", "content": "你是一个信息抽取助手"},
#                 {"role": "user", "content": prompt}
#             ]
#         )
#         return response.choices[0].message.content


# # =======================
# # 单条文本处理
# # =======================
# def weak_label(text, llm_client):
#     prompt = PROMPT_TEMPLATE.format(input_text=text)
#     response = llm_client.generate(prompt)
#     try:
#         data = json.loads(response)
#         return clean_output(data, RELATION_TYPES)
#     except Exception:
#         return {"entities": [], "relations": []}


# # =======================
# # 并发 pipeline
# # =======================
# def pipeline(texts, llm_client, max_workers=5):
#     results = []
#     with ThreadPoolExecutor(max_workers=max_workers) as executor:
#         future_to_text = {executor.submit(weak_label, t, llm_client): t for t in texts}
#         for future in tqdm(as_completed(future_to_text), total=len(texts), desc="Processing texts", unit="text"):
#             try:
#                 results.append(future.result())
#             except Exception as e:
#                 print("❌ 出错:", e)
#                 results.append({"entities": [], "relations": []})
#     return results


# # =======================
# # 主程序
# # =======================
# if __name__ == "__main__":
#     doc = Document("C:/PHD/2025/08/food_nutrition.docx")
#     texts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

#     llm_client = OpenAILLM()
#     results = pipeline(texts, llm_client, max_workers=5)  # ✅ 并发 5 个任务

#     with open("spert_dataset.jsonl", "w", encoding="utf-8") as f:
#         for item in results:
#             f.write(json.dumps(item, ensure_ascii=False) + "\n")

#     print("✅ spert_dataset.jsonl 已生成")

import json
from schema import RELATION_TYPES
from rules import clean_output
from prompt_template import PROMPT_TEMPLATE
import os
from openai import OpenAI
from tqdm import tqdm
from docx import Document
import re
from concurrent.futures import ThreadPoolExecutor, as_completed


# ==============================
# 代理设置（如有）
# ==============================
os.environ["HTTP_PROXY"] = "http://127.0.0.1:7890"
os.environ["HTTPS_PROXY"] = "http://127.0.0.1:7890"


# ==============================
# GPT 封装
# ==============================
class OpenAILLM:
    def __init__(self):
        self.client = OpenAI(api_key="")
    def generate(self, prompt):
        response = self.client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "你是一个信息抽取助手"},
                {"role": "user", "content": prompt}
            ]
        )
        return response.choices[0].message.content


# ==============================
# 判断标题类型
# ==============================
def is_main_chapter(text):
    return bool(re.match(r"^第[一二三四五六七八九十百]+章", text))

def is_subsection(text):
    return bool(re.match(r"^\d+\.\d+(\.\d+)?", text))  # 支持 1.2 或 1.2.1


# ==============================
# 从 Word 中提取结构化段落
# ==============================
def extract_structured_paragraphs(doc_path):
    doc = Document(doc_path)
    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

    current_main = ""
    current_sub = ""
    current_text = ""
    structured_samples = []

    for para in paragraphs:
        if is_main_chapter(para):
            current_main = para
            continue
        elif is_subsection(para):
            if current_text:
                structured_samples.append({
                    "main_section": current_main,
                    "subsection": current_sub,
                    "text": current_text.strip()
                })
            current_sub = para
            current_text = ""
        else:
            current_text += para + " "

    if current_text:
        structured_samples.append({
            "main_section": current_main,
            "subsection": current_sub,
            "text": current_text.strip()
        })

    return structured_samples


# ==============================
# 调用 GPT 弱标注
# ==============================
def weak_label(sample, llm_client):
    context = retrieve_context(sample["text"])  # 🔍 从向量库检索相关背景
    prompt = PROMPT_TEMPLATE.format(context=context, input_text=sample["text"])
    response = llm_client.generate(prompt)
    try:
        data = json.loads(response)
        result = clean_output(data, RELATION_TYPES)
        result["main_section"] = sample["main_section"]
        result["subsection"] = sample["subsection"]
        return result
    except Exception as e:
        print("❌ JSON 解析失败:", e)
        return {
            "main_section": sample["main_section"],
            "subsection": sample["subsection"],
            "entities": [],
            "relations": []
        }


# ==============================
# 并发 pipeline
# ==============================
def pipeline(samples, llm_client, max_workers=5):
    results = []
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {
            executor.submit(weak_label, sample, llm_client): sample for sample in samples
        }
        for future in tqdm(as_completed(futures), total=len(samples), desc="Processing sections", unit="section"):
            try:
                results.append(future.result())
            except Exception as e:
                sample = futures[future]
                print("❌ 执行失败:", e)
                results.append({
                    "main_section": sample["main_section"],
                    "subsection": sample["subsection"],
                    "entities": [],
                    "relations": []
                })
    return results


# ==============================
# 主程序
# ==============================
if __name__ == "__main__":
    doc_path = "C:/PHD/2025/08/food_nutrition.docx"  # ✅ 替换为你的 docx 路径

    print("📖 正在结构化提取章节段落...")
    samples = extract_structured_paragraphs(doc_path)
    print(f"✅ 共提取 {len(samples)} 条结构化段落")

    llm_client = OpenAILLM()
    results = pipeline(samples, llm_client, max_workers=5)

    print("💾 正在写入 spert_dataset.jsonl...")
    with open("spert_dataset.jsonl", "w", encoding="utf-8") as f:
        for item in results:
            # 删除结构字段
            item.pop("main_section", None)
            item.pop("subsection", None)
            f.write(json.dumps(item, ensure_ascii=False) + "\n")


    print("✅ 数据生成完毕，文件为：spert_dataset.jsonl")
